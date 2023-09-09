from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Create a new Word document
doc = Document()

# Add a title to the document
title = doc.add_heading('My Document', level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add some paragraphs to the document
for i in range(1, 6):
    paragraph = doc.add_paragraph(f'This is paragraph {i}.')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = paragraph.runs[0]
    run.bold = True
    run.font.size = Pt(12)

# Add a table to the document
table = doc.add_table(rows=3, cols=3)
table.autofit = False

# Set table width and alignment
table.width = Pt(400)
table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add content to the table cells
for row in table.rows:
    for cell in row.cells:
        cell.text = 'Cell Content'
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add a page break
page_break = OxmlElement('w:p')
run = page_break.makeelement('w:r')
br = page_break.makeelement('w:br')
run.append(br)
page_break.append(run)
doc._body.append(page_break)

# Add a new section with landscape orientation
section = doc.sections[-1]
section.orientation = WD_SECTION_ORIENTATION.LANDSCAPE

# Save the document
doc.save('my_document.docx')

print("Document created and saved.")

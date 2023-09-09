import requests
import csv
from docx import Document
from docx.shared import Inches
import os
from PIL import Image, UnidentifiedImageError
# Your Google Custom Search API key
api_key = 'AIzaSyBmkwEfKuManuqeVKCQNRkNZPOsekO8kyU'

# Your Custom Search Engine ID (CSE ID)
cse_id = '071af49c376f14afd'

# Query for laptops on Amazon
query = 'hp laptops'

image_dir = 'images'
os.makedirs(image_dir, exist_ok=True)

# Construct the API URL with the CSE ID
url = f'https://www.googleapis.com/customsearch/v1?key={api_key}&cx={cse_id}&q={query}'
# Send a GET request to the Google Custom Search API
response = requests.get(url)

# Check if the request was successful
if response.status_code == 200:
    data = response.json()
    if 'items' in data:
        laptop_results = data['items']
        if laptop_results:
            # Create a CSV file
            with open('laptops.csv', 'w', newline='', encoding='utf-8') as csvfile:
                csv_fieldnames = ['Laptop Name', 'Link', 'Image']
                csv_writer = csv.DictWriter(csvfile, fieldnames=csv_fieldnames)
                csv_writer.writeheader()
                
                # Create a Word document
                doc = Document()
                
                for result in laptop_results:
                    laptop_name = result.get('title', '').strip()
                    laptop_link = result.get('link', '')
                    laptop_image_url = result.get('pagemap', {}).get('cse_image', [{}])[0].get('src', '')
                    
                    # Download the image and save it locally
                    if laptop_image_url:
                        image_filename = os.path.join(image_dir, f'{laptop_name}.jpg')
                        try:
                            response_image = requests.get(laptop_image_url)
                            with open(image_filename, 'wb') as img_file:
                                img_file.write(response_image.content)
                            
                            # Open the downloaded image using Pillow (PIL)
                            image = Image.open(image_filename)
                            
                            # Convert the image to RGB mode if it's not already in RGB mode
                            if image.mode != 'RGB':
                                image = image.convert('RGB')
                            
                            # Specify the image format (e.g., JPEG) when saving
                            image.save(image_filename, format='JPEG')
                            
                        except UnidentifiedImageError:
                            # Handle the UnidentifiedImageError and skip this image
                            print(f"Skipping image: {laptop_name} (UnidentifiedImageError)")
                            continue
                    
                    # Add a blank line for separation in CSV
                    csv_writer.writerow({'Laptop Name': '', 'Link': '', 'Image': ''})
                    
                    # Add laptop attributes to CSV
                    csv_writer.writerow({'Laptop Name': laptop_name, 'Link': laptop_link, 'Image': image_filename})
                    
                    # Add a blank line for separation in the Word document
                    doc.add_paragraph('')
                    
                    # Add laptop attributes to the Word document
                    doc.add_heading('Laptop Name:', level=1)
                    doc.add_paragraph(laptop_name)
                    
                    doc.add_heading('Link:', level=1)
                    doc.add_paragraph(laptop_link)
                    
                    # Add the image to the Word document if available
                    if laptop_image_url:
                        doc.add_heading('Image:', level=1)
                        doc.add_picture(image_filename, width=Inches(4.0))
                
                # Save the CSV file
                print("Data has been successfully saved to 'laptops.csv'.")
                
                # Save the Word document
                doc.save('laptops.docx')
                
                print("Data has been successfully saved to 'laptops.docx'.")
        else:
            print("No search results found.")
    else:
        print("No 'items' found in the API response.")
else:
    print(f"Failed to retrieve data from the Google Custom Search API. Status Code: {response.status_code}")
    print(response.text) 